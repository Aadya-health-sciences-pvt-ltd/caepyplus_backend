"""Plain-text extraction for Office document resumes (.doc / .docx).

Gemini Vision only understands PDFs and images, so Word documents are first
converted to plain text here and then fed through the text-based extraction
path (``ResumeExtractionService.extract_from_text``).

- ``.docx`` is parsed in-process with ``python-docx`` (pure Python, no system
  dependency).
- ``.doc`` is the legacy OLE2 binary format, which has no reliable pure-Python
  parser. We still accept the upload and try the ``.docx`` parser (some files
  are ``.docx`` mislabelled as ``.doc``); a genuine binary ``.doc`` returns a
  clear "convert to PDF/.docx" message.
"""
from __future__ import annotations

import io
import logging

from ..core.exceptions import FileValidationError

logger = logging.getLogger(__name__)

# Extensions handled via text extraction (vs. Gemini Vision).
OFFICE_EXTENSIONS: frozenset[str] = frozenset({"doc", "docx"})


def _extract_docx(content: bytes) -> str:
    """Extract text from a .docx file (paragraphs + table cells)."""
    from docx import Document

    document = Document(io.BytesIO(content))

    parts: list[str] = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text and paragraph.text.strip()
    ]

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def _extract_doc(content: bytes) -> str:
    """Best-effort text extraction for a .doc upload.

    Legacy binary .doc has no reliable pure-Python parser. Some uploads are
    actually .docx renamed to .doc, so we try the .docx parser before failing
    with a user-facing message.
    """
    try:
        return _extract_docx(content)
    except Exception:  # noqa: BLE001 - best-effort fallback
        logger.info("Legacy .doc could not be parsed as a .docx fallback")

    raise FileValidationError(
        message="Unable to read this .doc file. Please upload a PDF or .docx instead.",
    )


def extract_text_from_document(content: bytes, extension: str) -> str:
    """Return plain text from a .doc/.docx resume.

    Args:
        content: Raw file bytes.
        extension: Lower-cased file extension without the dot (``doc``/``docx``).

    Raises:
        FileValidationError: If the type is unsupported or no text can be read.
    """
    ext = extension.lower()

    if ext == "docx":
        text = _extract_docx(content)
    elif ext == "doc":
        text = _extract_doc(content)
    else:
        raise FileValidationError(message=f"Unsupported document type: {ext}")

    if not text or not text.strip():
        raise FileValidationError(
            message="No readable text found in the document.",
        )

    return text
