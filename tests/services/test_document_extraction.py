"""Unit tests for Word document (.doc/.docx) text extraction."""
from __future__ import annotations

import io

import pytest

from src.app.core.exceptions import FileValidationError
from src.app.services.document_extraction import (
    OFFICE_EXTENSIONS,
    extract_text_from_document,
)


def _make_docx_bytes(paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    """Build an in-memory .docx file with the given paragraphs/table."""
    from docx import Document

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)

    if table:
        rows = len(table)
        cols = len(table[0])
        tbl = document.add_table(rows=rows, cols=cols)
        for r, row in enumerate(table):
            for c, value in enumerate(row):
                tbl.rows[r].cells[c].text = value

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_office_extensions_contains_doc_and_docx() -> None:
    assert set(OFFICE_EXTENSIONS) == {"doc", "docx"}


def test_extract_docx_paragraphs() -> None:
    content = _make_docx_bytes(["Dr. Asha Rao", "Cardiologist", "MBBS, MD"])

    text = extract_text_from_document(content, "docx")

    assert "Dr. Asha Rao" in text
    assert "Cardiologist" in text
    assert "MBBS, MD" in text


def test_extract_docx_includes_table_cells() -> None:
    content = _make_docx_bytes(
        ["Qualifications"],
        table=[["Degree", "Year"], ["MBBS", "2008"]],
    )

    text = extract_text_from_document(content, "docx")

    assert "Degree | Year" in text
    assert "MBBS | 2008" in text


def test_extract_docx_empty_raises() -> None:
    content = _make_docx_bytes([])

    with pytest.raises(FileValidationError):
        extract_text_from_document(content, "docx")


def test_unsupported_extension_raises() -> None:
    with pytest.raises(FileValidationError):
        extract_text_from_document(b"whatever", "rtf")


def test_doc_falls_back_to_docx_when_mislabelled() -> None:
    """A .docx renamed to .doc should still extract when antiword is absent."""
    content = _make_docx_bytes(["Mislabelled docx content"])

    text = extract_text_from_document(content, "doc")

    assert "Mislabelled docx content" in text


def test_doc_unreadable_raises() -> None:
    with pytest.raises(FileValidationError):
        extract_text_from_document(b"not a real doc file", "doc")
